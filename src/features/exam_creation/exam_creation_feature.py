# =============================================================================
# exam_creation_feature.py — Exam Creation Feature
# =============================================================================
# Provides the full Exam Creation UI and all supporting database operations.
#
# Feature overview:
#   Teachers generate exam questions in two modes:
#     Variation mode    — upload or paste existing questions and produce
#                         reworded variants at the same difficulty level.
#     From Topics mode  — upload or paste lecture content and generate a
#                         fresh question set at a chosen difficulty level.
#   Generated questions are displayed immediately, saved to the database,
#   and accessible through the History tab where they can be edited, downloaded,
#   or deleted.
#
# Database integration:
#   - Source files uploaded by the teacher are saved to disk under the
#     exam_creation subdirectory for the current assessment and recorded in the
#     files table with feature_name='exam_creation'. Saved files can be reused
#     across generation runs and deleted individually from within the feature.
#   - After a successful LLM generation, all questions from the run are saved
#     to exam_creation_questions, grouped by a creation_session_id UUID.
#   - The History tab queries sessions belonging to the current user and
#     assessment, and allows per-question editing and per-session deletion.
#
# Course/assessment context:
#   This feature is always rendered inside a specific assessment. The selected
#   course and assessment are read from session state using the tab-namespaced
#   keys set by app.py's navigation system:
#     st.session_state["exam_creation_selected_course"]     -> {"id": int, "name": str}
#     st.session_state["exam_creation_selected_assessment"] -> {"id": int, "title": str}
# =============================================================================

import io
import os
import uuid
import json
import streamlit as st
import pandas as pd
from pathlib import Path
from typing import List, Dict, Any
from sqlalchemy import text
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from db import get_connection, get_engine
from auth import save_uploaded_file, delete_physical_file

from src.utils.llm_utils import MODELS, generate_llm_response, strip_llm_json
from src.utils.pdf_utils import extract_text_from_pdf, save_uploaded_pdf


# =============================================================================
# DATABASE WRITE OPERATIONS
# =============================================================================

def save_exam_creation_session(
    creation_session_id: str,
    user_id: int,
    assessment_id: int,
    creation_mode: str,
    questions: List[Dict],
) -> None:
    """
    Persist all questions from one generation run to exam_creation_questions.

    One row is inserted per question. All rows share the same creation_session_id
    so the History tab can group and display them as a single session.

    creation_mode must be either 'variation' or 'from_topics', matching the
    ENUM defined in the schema.

    Each question dict is expected to contain:
        question_text        (str, required)
        answer_guidance      (str or None)
        original_question_text (str or None — variation mode only)
        question_type        (str or None — from_topics mode only)
        topic                (str or None — from_topics mode only)
        difficulty           (str or None — from_topics mode only)

    assessment_id links the session to the assessment the teacher was working
    under, scoping the History tab to the current assessment.
    """
    conn = get_connection()
    cursor = conn.cursor()
    try:
        for q in questions:
            cursor.execute("""
                INSERT INTO exam_creation_questions (
                    creation_session_id,
                    user_id,
                    assessment_id,
                    creation_mode,
                    original_question_text,
                    question_text,
                    question_type,
                    topic,
                    difficulty,
                    answer_guidance
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (
                creation_session_id,
                user_id,
                assessment_id,
                creation_mode,
                q.get("original_question_text"),
                q.get("question_text", ""),
                q.get("question_type"),
                q.get("topic"),
                q.get("difficulty"),
                q.get("answer_guidance"),
            ))
        conn.commit()
    finally:
        cursor.close()
        conn.close()


def update_exam_creation_question(
    question_id: int,
    user_id: int,
    question_text: str,
    answer_guidance: str,
) -> None:
    """
    Update the question text and answer guidance for a single question row.

    Filtered by both question id and user_id so teachers can only edit their
    own questions. Called from the History tab edit controls.
    """
    conn = get_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("""
            UPDATE exam_creation_questions
            SET question_text   = %s,
                answer_guidance = %s
            WHERE id      = %s
              AND user_id = %s
        """, (question_text, answer_guidance, question_id, user_id))
        conn.commit()
    finally:
        cursor.close()
        conn.close()


def delete_exam_creation_session(creation_session_id: str, user_id: int) -> None:
    """
    Delete all question rows belonging to a creation session.

    Filtered by both creation_session_id and user_id so teachers can only
    delete their own sessions. Called from the History tab delete controls.
    """
    conn = get_connection()
    cursor = conn.cursor()
    try:
        cursor.execute("""
            DELETE FROM exam_creation_questions
            WHERE creation_session_id = %s
              AND user_id = %s
        """, (creation_session_id, user_id))
        conn.commit()
    finally:
        cursor.close()
        conn.close()


# =============================================================================
# DATABASE READ OPERATIONS
# =============================================================================

def get_exam_creation_sessions(user_id: int, assessment_id: int) -> List[Dict]:
    """
    Return all creation sessions for the given user and assessment.

    Each row represents one generation run, grouped by creation_session_id.
    Ordered newest first so the most recent session appears at the top of
    the History tab.
    """
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute("""
            SELECT
                creation_session_id,
                creation_mode,
                MAX(created_at)  AS created_at,
                COUNT(*)         AS question_count
            FROM exam_creation_questions
            WHERE user_id      = %s
              AND assessment_id = %s
            GROUP BY creation_session_id, creation_mode
            ORDER BY created_at DESC
        """, (user_id, assessment_id))
        return cursor.fetchall() or []
    finally:
        cursor.close()
        conn.close()


def get_exam_creation_session_questions(
    creation_session_id: str, user_id: int
) -> List[Dict]:
    """
    Return all question rows belonging to a specific creation session.

    Filtered by user_id so teachers can only read their own questions.
    Ordered by row id (insertion order) so questions are displayed in the
    sequence they were generated.
    """
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute("""
            SELECT
                id,
                question_text,
                question_type,
                topic,
                difficulty,
                answer_guidance,
                original_question_text,
                creation_mode
            FROM exam_creation_questions
            WHERE creation_session_id = %s
              AND user_id = %s
            ORDER BY id ASC
        """, (creation_session_id, user_id))
        return cursor.fetchall() or []
    finally:
        cursor.close()
        conn.close()


# =============================================================================
# FILE DB OPERATIONS
# =============================================================================

def get_exam_creation_files(assessment_id: int) -> List[Dict]:
    """
    Return all files saved under the exam_creation feature for the given assessment.

    Filters by both assessment_id and feature_name so only files uploaded
    through the Exam Creation feature are returned, not general course files
    or files from other features within the same assessment.

    Ordered newest first so recently uploaded files appear at the top of
    the saved files list.
    """
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        cursor.execute("""
            SELECT id, file_name, file_path, uploaded_at
            FROM files
            WHERE assessment_id = %s
              AND feature_name  = 'exam_creation'
            ORDER BY uploaded_at DESC
        """, (assessment_id,))
        return cursor.fetchall() or []
    finally:
        cursor.close()
        conn.close()


def save_exam_creation_file(
    file_bytes: bytes,
    original_name: str,
    course_id: int,
    course_name: str,
    assessment_id: int,
    assessment_title: str,
    user_id: int,
) -> None:
    """
    Save an uploaded source file to the exam_creation directory for the
    current assessment and insert a record into the files table.

    Uses the shared save_uploaded_file() helper from auth.py, which handles
    directory creation and UUID-prefixed naming to prevent collisions.
    The feature_name column is set to 'exam_creation' so the file appears
    only in the Exam Creation saved files list, not in the general Files view
    for other features.

    Errors are surfaced via st.warning rather than raising, so that a storage
    failure does not block the text extraction workflow from continuing.
    """
    try:
        saved_name, saved_path = save_uploaded_file(
            file_bytes=file_bytes,
            original_name=original_name,
            course_name=course_name,
            assessment_name=assessment_title,
            course_id=course_id,
            feature_name="exam_creation",
        )
        engine = get_engine()
        with engine.begin() as conn:
            conn.execute(text("""
                INSERT INTO files
                    (file_name, file_path, course_id, assessment_id, uploaded_by, feature_name)
                VALUES (:name, :path, :cid, :aid, :uid, :feat)
            """), {
                "name": saved_name,
                "path": saved_path,
                "cid":  course_id,
                "aid":  assessment_id,
                "uid":  user_id,
                "feat": "exam_creation",
            })
    except Exception as exc:
        st.warning(f"File '{original_name}' could not be saved: {exc}")


# =============================================================================
# LLM RESPONSE PARSERS
# =============================================================================

def _parse_variation_response(variations_json: str) -> List[Dict]:
    """
    Parse a variation-mode LLM response into a flat list of question dicts
    suitable for passing to save_exam_creation_session().

    The LLM returns a nested structure where each original question has a list
    of variations. This flattens that structure so each variation becomes one
    row in exam_creation_questions, with original_question_text storing the
    source question for reference.

    Returns an empty list if parsing fails, which causes the save step to be
    skipped without interrupting the display flow.
    """
    try:
        data = json.loads(strip_llm_json(variations_json))
        questions = []
        for question_set in data.get("variations", []):
            original = question_set.get("original_question", "")
            for variation in question_set.get("variations", []):
                questions.append({
                    "original_question_text": original,
                    "question_text":          variation.get("question_text", ""),
                    "answer_guidance":        variation.get("answer_guidance"),
                    "question_type":          None,
                    "topic":                  None,
                    "difficulty":             None,
                })
        return questions
    except Exception:
        return []


def _parse_from_topics_response(exam_json: str) -> List[Dict]:
    """
    Parse a from-topics-mode LLM response into a flat list of question dicts
    suitable for passing to save_exam_creation_session().

    The LLM returns a flat list of question objects. Each is mapped directly
    to a row in exam_creation_questions. original_question_text is None for
    this mode as there is no source question being varied.

    Returns an empty list if parsing fails, which causes the save step to be
    skipped without interrupting the display flow.
    """
    try:
        data = json.loads(strip_llm_json(exam_json))
        questions = []
        for q in data.get("questions", []):
            questions.append({
                "original_question_text": None,
                "question_text":          q.get("question_text", ""),
                "question_type":          q.get("question_type"),
                "topic":                  q.get("topic"),
                "difficulty":             q.get("difficulty"),
                "answer_guidance":        q.get("answer_guidance"),
            })
        return questions
    except Exception:
        return []


# =============================================================================
# LLM PROMPT BUILDERS
# =============================================================================

def create_question_variation_prompt(
    original_questions: str, num_variations: int = 3
) -> str:
    """
    Build the LLM prompt for variation mode.

    Instructs the model to generate num_variations reworded versions of each
    input question at the same difficulty level and testing the same concepts.
    The response must be a JSON object matching the variations schema so that
    display_question_variations() and _parse_variation_response() can process it.
    """
    return f"""You are an expert education professional tasked with creating variations of exam questions.
I will provide you with original exam questions or lecture topics, and your task is to create {num_variations} variations
for each question while maintaining the same difficulty level and testing the same concepts.

GUIDELINES:
1. Preserve the original learning objectives and complexity
2. Create questions that test the same knowledge but use different scenarios or phrasing
3. Include a mix of question types (multiple choice, short answer, problem-solving)
4. Format your response as a well-structured JSON object

ORIGINAL QUESTIONS/TOPICS:
{original_questions}

Generate {num_variations} variations for each question/topic and output in the following JSON format:
{{
  "variations": [
    {{
      "original_question": "The text of the original question",
      "variations": [
        {{
          "variation_number": 1,
          "question_text": "First variation of the question",
          "answer_guidance": "Guidance for what would constitute a good answer"
        }},
        {{
          "variation_number": 2,
          "question_text": "Second variation of the question",
          "answer_guidance": "Guidance for what would constitute a good answer"
        }},
        ...
      ]
    }},
    ...
  ]
}}
"""


def create_exam_from_topics_prompt(
    topics: str, num_questions: int, difficulty_level: str
) -> str:
    """
    Build the LLM prompt for from-topics mode.

    Instructs the model to generate num_questions new questions at the given
    difficulty level based on the provided lecture content. The response must
    be a JSON object matching the questions schema so that
    display_generated_exam() and _parse_from_topics_response() can process it.
    """
    return f"""You are an expert education professional tasked with creating exam questions.
I will provide you with lecture topics/content, and your task is to create {num_questions} exam questions
at a {difficulty_level} difficulty level that thoroughly assess understanding of these topics.

LECTURE TOPICS/CONTENT:
{topics}

GUIDELINES:
1. Create questions that assess critical thinking and comprehension
2. Include a mix of question types (multiple choice, short answer, problem-solving)
3. Make sure questions cover the full range of topics provided
4. For each question, provide answer guidance for the instructor
5. Format your response as a well-structured JSON object

Generate {num_questions} questions at {difficulty_level} difficulty and output in the following JSON format:
{{
  "questions": [
    {{
      "question_number": 1,
      "question_text": "The text of the question",
      "question_type": "multiple-choice|short-answer|problem-solving",
      "topic": "The specific topic this question addresses",
      "difficulty": "{difficulty_level}",
      "answer_guidance": "Guidance for what would constitute a good answer"
    }},
    ...
  ]
}}
"""


# =============================================================================
# DISPLAY HELPERS
# =============================================================================

def display_question_variations(variations_json: str) -> None:
    """
    Render a variation-mode LLM response in the UI.

    Displays each original question followed by its numbered variants in
    collapsible expanders. Shows the raw response as a code block if parsing
    fails so the teacher can inspect what the model returned.
    """
    try:
        data = json.loads(strip_llm_json(variations_json))
        if "variations" in data:
            for question_set in data["variations"]:
                original   = question_set.get("original_question", "")
                variations = question_set.get("variations", [])

                st.markdown("### Original Question")
                st.write(original)

                st.markdown(f"### Variations ({len(variations)})")
                for i, variation in enumerate(variations):
                    variation_number = variation.get("variation_number", i + 1)
                    with st.expander(f"Variation {variation_number}"):
                        st.markdown("**Question:**")
                        st.write(variation.get("question_text", ""))
                        st.markdown("**Answer Guidance:**")
                        st.write(variation.get("answer_guidance", ""))

                st.divider()
    except Exception as e:
        st.error(f"Error parsing response: {str(e)}")
        st.write("Raw response:")
        st.code(variations_json)


def display_generated_exam(exam_json: str) -> None:
    """
    Render a from-topics-mode LLM response in the UI.

    Displays each generated question in a collapsible expander showing the
    question text, type, topic, difficulty, and answer guidance. Shows the
    raw response as a code block if parsing fails.
    """
    try:
        data = json.loads(strip_llm_json(exam_json))
        if "questions" in data:
            questions = data["questions"]
            st.markdown(f"### Generated Exam ({len(questions)} questions)")
            for question in questions:
                question_number = question.get("question_number", "")
                question_text   = question.get("question_text", "")
                question_type   = question.get("question_type", "")
                topic           = question.get("topic", "")
                difficulty      = question.get("difficulty", "")
                answer_guidance = question.get("answer_guidance", "")

                with st.expander(f"Question {question_number} - {topic} ({question_type})"):
                    st.markdown("**Question:**")
                    st.write(question_text)
                    st.markdown(f"**Topic:** {topic}")
                    st.markdown(f"**Type:** {question_type}")
                    st.markdown(f"**Difficulty:** {difficulty}")
                    st.markdown("**Answer Guidance:**")
                    st.write(answer_guidance)
    except Exception as e:
        st.error(f"Error parsing response: {str(e)}")
        st.write("Raw response:")
        st.code(exam_json)


def _create_exam_word_document(
    questions: List[Dict],
    creation_mode: str,
    title: str = "Generated Exam Questions",
) -> io.BytesIO:
    """
    Build a Word document from a list of exam creation question dicts and
    return it as a BytesIO object ready for st.download_button.

    Handles both creation modes in a single pass:
      variation   — each question may carry original_question_text showing
                    the source question it was derived from. When present,
                    the original is printed as an italic caption above the
                    generated question so the teacher has full context.
      from_topics — each question may carry question_type, topic, and
                    difficulty metadata. When present, these are printed as
                    an italic caption below the answer guidance.

    Fields that are absent or None are skipped silently so the function
    works correctly for partial records and for both modes without branching.

    Args:
        questions:      List of question dicts from _parse_variation_response(),
                        _parse_from_topics_response(), or
                        get_exam_creation_session_questions().
        creation_mode:  'variation' or 'from_topics'. Used only to set the
                        document subtitle; field rendering is driven by the
                        presence of each field in the question dicts.
        title:          Document heading. Defaults to a generic title; callers
                        pass a more specific label where available.

    Returns:
        io.BytesIO positioned at offset 0, ready to be passed directly to
        st.download_button's data parameter via .getvalue().

    Raises:
        Exception propagated from python-docx if document creation fails.
        Callers wrap this function in try/except and show st.error().
    """
    doc = Document()

    # ── Title and subtitle ────────────────────────────────────────────────────
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    mode_label = "Question Variations" if creation_mode == "variation" else "Exam from Topics"
    subtitle = doc.add_paragraph(f"Mode: {mode_label}  ·  {len(questions)} question(s)")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # ── Questions ─────────────────────────────────────────────────────────────
    for i, q in enumerate(questions, start=1):
        question_text   = q.get("question_text", "")
        answer_guidance = q.get("answer_guidance") or ""
        original        = q.get("original_question_text") or ""
        question_type   = q.get("question_type") or ""
        topic           = q.get("topic") or ""
        difficulty      = q.get("difficulty") or ""

        # For variation-mode questions, print the source question above the
        # generated variant so the teacher can see what it was derived from.
        if original:
            orig_para = doc.add_paragraph()
            orig_run  = orig_para.add_run(f"Original: {original}")
            orig_run.italic = True

        # Question text — bold numbered label followed by plain question body.
        q_para = doc.add_paragraph()
        q_para.add_run(f"Question {i}: ").bold = True
        q_para.add_run(question_text)

        # Answer guidance — always included since this document is for the
        # instructor. Empty guidance is written as an em dash placeholder.
        guidance_para = doc.add_paragraph()
        guidance_para.add_run("Answer Guidance: ").bold = True
        guidance_para.add_run(answer_guidance if answer_guidance else "\u2014")

        # For from-topics questions, print the metadata line below the guidance
        # so the teacher can see type, topic, and difficulty at a glance.
        if question_type or topic or difficulty:
            meta_parts = []
            if question_type:
                meta_parts.append(f"Type: {question_type}")
            if topic:
                meta_parts.append(f"Topic: {topic}")
            if difficulty:
                meta_parts.append(f"Difficulty: {difficulty}")
            meta_para = doc.add_paragraph("  ·  ".join(meta_parts))
            meta_para.runs[0].italic = True

        doc.add_paragraph("")

    # ── Serialise to BytesIO ──────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# =============================================================================
# LLM GENERATION WRAPPERS
# =============================================================================

def generate_question_variations(
    original_questions: str, model_name: str, num_variations: int = 3
) -> str:
    """Send the variation prompt to the LLM and return the raw response string."""
    prompt = create_question_variation_prompt(original_questions, num_variations)
    return generate_llm_response(prompt, model_name)


def generate_exam_from_topics(
    topics: str, model_name: str, num_questions: int, difficulty_level: str
) -> str:
    """Send the from-topics prompt to the LLM and return the raw response string."""
    prompt = create_exam_from_topics_prompt(topics, num_questions, difficulty_level)
    return generate_llm_response(prompt, model_name)


# =============================================================================
# MAIN UI
# =============================================================================

@st.dialog("Delete Session")
def _dialog_delete_exam_creation_session(session_id: str, user_id: int) -> None:
    """
    Confirmation modal for permanently deleting an exam creation session.

    Deletes all question rows belonging to the session from the database.
    """
    st.warning("Are you sure you want to delete this session? All generated questions will be permanently removed.")
    col1, col2 = st.columns(2)
    if col1.button("Delete", type="primary", key="ec_dialog_confirm_delete"):
        delete_exam_creation_session(session_id, user_id)
        st.toast("Session deleted.")
        st.rerun()
    if col2.button("Cancel", key="ec_dialog_cancel_delete"):
        st.rerun()


@st.dialog("Delete File")
def _dialog_delete_ec_file(file_id: int, file_name: str, file_path: str) -> None:
    """
    Confirmation modal for permanently deleting a saved exam creation source file.

    Removes the database record and the physical file from disk. The file_path
    is the relative path stored in the files table; deletion is handled by
    delete_physical_file() which catches errors silently so a missing file on
    disk does not block the DB record from being removed.
    """
    st.warning(
        f"Are you sure you want to delete **{file_name}**? "
        "This cannot be undone."
    )
    col1, col2 = st.columns(2)
    if col1.button("Delete", type="primary", key="ec_file_dialog_confirm"):
        conn = get_connection()
        cursor = conn.cursor()
        try:
            cursor.execute("DELETE FROM files WHERE id = %s", (file_id,))
            conn.commit()
        finally:
            cursor.close()
            conn.close()
        delete_physical_file(file_path)
        st.toast(f"File deleted.")
        st.rerun()
    if col2.button("Cancel", key="ec_file_dialog_cancel"):
        st.rerun()


def exam_creation_ui() -> None:
    """
    Render the full Exam Creation feature UI.

    The feature is always opened from within a specific course assessment. The
    current course and assessment are read from tab-namespaced session state
    keys so that generated questions are linked to the correct assessment.

    Tab structure:
        Generate Question Variations -- produce reworded variants of existing questions.
        Create Exam from Topics      -- generate a fresh question set from lecture content.
        History                      -- browse, edit, download, and delete past sessions.
    """
    st.subheader("Exam Creation Assistant")

    # Read the course and assessment context set by app.py's navigation.
    # These values are used when saving generated questions and source files
    # to the database and to disk.
    _course       = st.session_state.get("exam_creation_selected_course", {}) or {}
    _assessment   = st.session_state.get("exam_creation_selected_assessment", {}) or {}
    course_id       = _course.get("id")
    course_name     = _course.get("name", "")
    assessment_id   = _assessment.get("id")
    assessment_title = _assessment.get("title", "")
    user_id         = st.session_state["user"]["id"]

    # Initialise the model preference key if not already set by
    # _load_model_preferences() at login.
    if "exam_creation_selected_model" not in st.session_state:
        st.session_state.exam_creation_selected_model = list(MODELS.keys())[0]

    tab1, tab2, tab3 = st.tabs([
        "📝 Generate Question Variations",
        "📚 Create Exam from Topics",
        "🕘 History",
    ])

    # -------------------------------------------------------------------------
    # TAB 1 — Generate Question Variations
    # -------------------------------------------------------------------------
    with tab1:
        st.write("Upload or enter original exam questions to generate variations")

        input_method = st.radio(
            "How would you like to input the original questions?",
            ["Upload new file", "Use saved file", "Enter questions manually"],
            key="ec_var_input_method",
        )

        original_questions = ""

        if input_method == "Upload new file":
            # Upload a new PDF, save it to the course directory under the
            # exam_creation subdirectory, extract its text, and use it as input.
            uploaded_file = st.file_uploader(
                "Upload a PDF containing the original questions",
                type=["pdf"],
                key="ec_var_uploader",
            )
            if uploaded_file:
                try:
                    # Guard both file persistence and text extraction behind
                    # a session state flag keyed on the uploaded filename.
                    # Streamlit rerenders this block on every interaction
                    # while a file is present in the uploader. Without this
                    # guard, file saving would create duplicates on disk and
                    # in the DB, and re-extraction would trigger the spinner
                    # on every widget interaction in the tab.
                    if st.session_state.get("ec_var_saved_filename") != uploaded_file.name:
                        with st.spinner("Extracting text from PDF..."):
                            if course_id and assessment_id:
                                save_exam_creation_file(
                                    file_bytes=uploaded_file.getvalue(),
                                    original_name=uploaded_file.name,
                                    course_id=course_id,
                                    course_name=course_name,
                                    assessment_id=assessment_id,
                                    assessment_title=assessment_title,
                                    user_id=user_id,
                                )
                            file_path = save_uploaded_pdf(uploaded_file)
                            extracted = extract_text_from_pdf(file_path)
                            if os.path.exists(file_path):
                                os.remove(file_path)
                            st.session_state["ec_var_saved_filename"] = uploaded_file.name
                            st.session_state["ec_var_extracted_text"] = extracted
                    original_questions = st.session_state.get("ec_var_extracted_text", "")
                except Exception as e:
                    st.error(f"Error processing file: {str(e)}")

        elif input_method == "Use saved file":
            # List all files previously saved under exam_creation for this
            # assessment. The teacher selects one and its text is extracted
            # from disk for use as the generation input.
            saved_files = get_exam_creation_files(assessment_id) if assessment_id else []
            if not saved_files:
                st.info("No saved files for this assessment yet. Upload a new file to get started.")
            else:
                file_options = {f["file_name"]: f for f in saved_files}
                selected_name = st.selectbox(
                    "Select a saved file:",
                    list(file_options.keys()),
                    key="ec_var_saved_select",
                )
                selected_file = file_options[selected_name]
                if st.button("Use this file", key="ec_var_use_saved"):
                    try:
                        file_path = selected_file["file_path"]
                        original_questions = extract_text_from_pdf(file_path)
                        st.session_state["ec_var_extracted_text"] = original_questions
                        st.success(f"Text extracted from '{selected_name}'.")
                    except Exception as e:
                        st.error(f"Error reading file: {str(e)}")

            # Retrieve previously extracted text across rerenders so the
            # teacher can proceed to generation without re-clicking Use.
            if not original_questions:
                original_questions = st.session_state.get("ec_var_extracted_text", "")

        else:
            original_questions = st.text_area(
                "Enter the original exam questions:",
                height=200,
                placeholder="Enter each question on a new line or paragraph...",
                key="ec_var_text_area",
            )
            # Clear any previously extracted text when switching to manual entry.
            st.session_state.pop("ec_var_extracted_text", None)

        # ── Saved Files Management ─────────────────────────────────────────
        # Always visible regardless of input method so teachers can manage
        # their uploaded files at any time within this tab.
        if assessment_id:
            with st.expander("📁 Manage Saved Files", expanded=False):
                saved_files_list = get_exam_creation_files(assessment_id)
                if not saved_files_list:
                    st.info("No files saved for this assessment yet.")
                else:
                    for f in saved_files_list:
                        col_name, col_date, col_del = st.columns([4, 2, 1])
                        with col_name:
                            st.markdown(f"`{f['file_name']}`")
                        with col_date:
                            st.caption(str(f["uploaded_at"]))
                        with col_del:
                            if st.button(
                                "Delete",
                                key=f"ec_var_del_file_{f['id']}",
                                type="primary",
                                width="stretch",
                            ):
                                _dialog_delete_ec_file(f["id"], f["file_name"], f["file_path"])

        if original_questions:
            with st.expander("Review Original Questions"):
                st.write(original_questions)

            num_variations = st.slider(
                "Number of variations to generate for each question:",
                min_value=1, max_value=5, value=3,
            )

            # Model selection — honours the user's saved preference loaded at
            # login by _load_model_preferences() into exam_creation_selected_model.
            # A separate widget key prevents Streamlit from overwriting the
            # preference on first render.
            model_keys  = list(MODELS.keys())
            saved_model = st.session_state.get("exam_creation_selected_model", model_keys[0])
            if saved_model not in model_keys:
                saved_model = model_keys[0]
            selected_model_key = st.selectbox(
                "Select the model to use:",
                model_keys,
                index=model_keys.index(saved_model),
                key="exam_creation_variation_model",
            )
            st.session_state.exam_creation_selected_model = selected_model_key
            selected_model = MODELS[selected_model_key]

            if selected_model == "llama-3.3-70b-groq" and not st.session_state.get("groq_api_key"):
                st.warning("⚠️ Groq API key is required. Please add your API key in your profile settings.")
            if selected_model == "gemini-2.5-flash" and not st.session_state.get("gemini_api_key"):
                st.warning("⚠️ Gemini API key is required. Please add your API key in your profile settings.")

            if st.button("Generate Question Variations"):
                with st.spinner("Generating question variations..."):
                    variations_json = generate_question_variations(
                        original_questions=original_questions,
                        model_name=selected_model,
                        num_variations=num_variations,
                    )
                    st.session_state.variations_response = variations_json

                    # Parse and save the generated questions to the database.
                    # Each variation becomes one row in exam_creation_questions,
                    # grouped under a new creation_session_id for this run.
                    parsed = _parse_variation_response(variations_json)
                    if parsed and assessment_id:
                        save_exam_creation_session(
                            creation_session_id=str(uuid.uuid4()),
                            user_id=user_id,
                            assessment_id=assessment_id,
                            creation_mode="variation",
                            questions=parsed,
                        )

                if hasattr(st.session_state, "variations_response"):
                    display_question_variations(st.session_state.variations_response)
                    # Parse and re-serialise the raw LLM response before download
                    # so the file is guaranteed to be clean, valid JSON regardless
                    # of any fencing or preamble in the original model output.
                    try:
                        _var_download_data = json.dumps(
                            json.loads(strip_llm_json(st.session_state.variations_response)),
                            indent=2,
                        )
                    except Exception:
                        _var_download_data = st.session_state.variations_response

                    # Explicit keys are required on both download buttons because
                    # Streamlit renders all tab content simultaneously. Without
                    # explicit keys, auto-generated keys are position-based in the
                    # render tree and can conflict with the equivalent buttons in
                    # Tab 2 which occupy the same structural position.
                    col_dl_json, col_dl_word = st.columns(2)
                    with col_dl_json:
                        st.download_button(
                            label="Download Variations as JSON",
                            data=_var_download_data,
                            file_name="question_variations.json",
                            mime="application/json",
                            key="ec_var_dl_json",
                        )
                    with col_dl_word:
                        try:
                            _var_word_questions = _parse_variation_response(
                                st.session_state.variations_response
                            )
                            # _parse_variation_response returns [] silently on any
                            # parse failure. Guard against this so the teacher is
                            # not given a download button that produces an empty
                            # document with no explanation.
                            if not _var_word_questions:
                                st.warning("Word export is unavailable — the response could not be parsed.")
                            else:
                                _var_word_buf = _create_exam_word_document(
                                    questions=_var_word_questions,
                                    creation_mode="variation",
                                    title="Question Variations",
                                )
                                st.download_button(
                                    label="Download Variations as Word",
                                    data=_var_word_buf.getvalue(),
                                    file_name="question_variations.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key="ec_var_dl_word",
                                )
                        except Exception as e:
                            st.error(f"Could not generate Word document: {e}")

    # -------------------------------------------------------------------------
    # TAB 2 — Create Exam from Topics
    # -------------------------------------------------------------------------
    with tab2:
        st.write("Generate new exam questions based on lecture topics")

        input_method = st.radio(
            "How would you like to input the lecture topics?",
            ["Upload new file", "Use saved file", "Enter topics manually"],
            key="ec_topics_input_method",
        )

        lecture_topics = ""

        if input_method == "Upload new file":
            # Upload a new PDF, save it to the course directory under the
            # exam_creation subdirectory, extract its text, and use it as input.
            uploaded_file = st.file_uploader(
                "Upload a PDF containing the lecture topics/content",
                type=["pdf"],
                key="ec_topics_uploader",
            )
            if uploaded_file:
                try:
                    # Guard both file persistence and text extraction behind
                    # a session state flag keyed on the uploaded filename.
                    # Streamlit rerenders this block on every interaction
                    # while a file is present in the uploader. Without this
                    # guard, file saving would create duplicates on disk and
                    # in the DB, and re-extraction would trigger the spinner
                    # on every widget interaction in the tab.
                    if st.session_state.get("ec_topics_saved_filename") != uploaded_file.name:
                        with st.spinner("Extracting text from PDF..."):
                            if course_id and assessment_id:
                                save_exam_creation_file(
                                    file_bytes=uploaded_file.getvalue(),
                                    original_name=uploaded_file.name,
                                    course_id=course_id,
                                    course_name=course_name,
                                    assessment_id=assessment_id,
                                    assessment_title=assessment_title,
                                    user_id=user_id,
                                )
                            file_path = save_uploaded_pdf(uploaded_file)
                            extracted = extract_text_from_pdf(file_path)
                            if os.path.exists(file_path):
                                os.remove(file_path)
                            st.session_state["ec_topics_saved_filename"] = uploaded_file.name
                            st.session_state["ec_topics_extracted_text"] = extracted
                    lecture_topics = st.session_state.get("ec_topics_extracted_text", "")
                except Exception as e:
                    st.error(f"Error processing file: {str(e)}")

        elif input_method == "Use saved file":
            # List all files previously saved under exam_creation for this
            # assessment. The teacher selects one and its text is extracted
            # from disk for use as the generation input.
            saved_files = get_exam_creation_files(assessment_id) if assessment_id else []
            if not saved_files:
                st.info("No saved files for this assessment yet. Upload a new file to get started.")
            else:
                file_options = {f["file_name"]: f for f in saved_files}
                selected_name = st.selectbox(
                    "Select a saved file:",
                    list(file_options.keys()),
                    key="ec_topics_saved_select",
                )
                selected_file = file_options[selected_name]
                if st.button("Use this file", key="ec_topics_use_saved"):
                    try:
                        file_path = selected_file["file_path"]
                        lecture_topics = extract_text_from_pdf(file_path)
                        st.session_state["ec_topics_extracted_text"] = lecture_topics
                        st.success(f"Text extracted from '{selected_name}'.")
                    except Exception as e:
                        st.error(f"Error reading file: {str(e)}")

            # Retrieve previously extracted text across rerenders so the
            # teacher can proceed to generation without re-clicking Use.
            if not lecture_topics:
                lecture_topics = st.session_state.get("ec_topics_extracted_text", "")

        else:
            lecture_topics = st.text_area(
                "Enter the lecture topics/content:",
                height=200,
                placeholder="Enter topics, key concepts, or lecture content...",
                key="ec_topics_text_area",
            )
            # Clear any previously extracted text when switching to manual entry.
            st.session_state.pop("ec_topics_extracted_text", None)

        # ── Saved Files Management ─────────────────────────────────────────
        # Always visible regardless of input method so teachers can manage
        # their uploaded files at any time within this tab.
        if assessment_id:
            with st.expander("📁 Manage Saved Files", expanded=False):
                saved_files_list = get_exam_creation_files(assessment_id)
                if not saved_files_list:
                    st.info("No files saved for this assessment yet.")
                else:
                    for f in saved_files_list:
                        col_name, col_date, col_del = st.columns([4, 2, 1])
                        with col_name:
                            st.markdown(f"`{f['file_name']}`")
                        with col_date:
                            st.caption(str(f["uploaded_at"]))
                        with col_del:
                            if st.button(
                                "Delete",
                                key=f"ec_topics_del_file_{f['id']}",
                                type="primary",
                                width="stretch",
                            ):
                                _dialog_delete_ec_file(f["id"], f["file_name"], f["file_path"])

        if lecture_topics:
            with st.expander("Review Lecture Topics"):
                st.write(lecture_topics)

            col1, col2 = st.columns(2)
            with col1:
                num_questions = st.slider(
                    "Number of questions to generate:",
                    min_value=5, max_value=30, value=10,
                )
            with col2:
                difficulty_level = st.select_slider(
                    "Difficulty level:",
                    options=["Easy", "Medium", "Hard"],
                    value="Medium",
                )

            # Model selection — reads the same exam_creation_selected_model key
            # written by the variation tab selectbox above (or set at login).
            # Both generation tabs share a single model preference so switching
            # tabs does not reset the model choice.
            model_keys  = list(MODELS.keys())
            saved_model = st.session_state.get("exam_creation_selected_model", model_keys[0])
            if saved_model not in model_keys:
                saved_model = model_keys[0]
            selected_model_key = st.selectbox(
                "Select the model to use:",
                model_keys,
                index=model_keys.index(saved_model),
                key="exam_creation_topics_model",
            )
            st.session_state.exam_creation_selected_model = selected_model_key
            selected_model = MODELS[selected_model_key]

            if selected_model == "llama-3.3-70b-groq" and not st.session_state.get("groq_api_key"):
                st.warning("⚠️ Groq API key is required. Please add your API key in your profile settings.")
            if selected_model == "gemini-2.5-flash" and not st.session_state.get("gemini_api_key"):
                st.warning("⚠️ Gemini API key is required. Please add your API key in your profile settings.")

            if st.button("Generate Exam Questions"):
                with st.spinner("Generating exam questions..."):
                    exam_json = generate_exam_from_topics(
                        topics=lecture_topics,
                        model_name=selected_model,
                        num_questions=num_questions,
                        difficulty_level=difficulty_level,
                    )
                    st.session_state.exam_response = exam_json

                    # Parse and save the generated questions to the database.
                    # Each question becomes one row in exam_creation_questions,
                    # grouped under a new creation_session_id for this run.
                    parsed = _parse_from_topics_response(exam_json)
                    if parsed and assessment_id:
                        save_exam_creation_session(
                            creation_session_id=str(uuid.uuid4()),
                            user_id=user_id,
                            assessment_id=assessment_id,
                            creation_mode="from_topics",
                            questions=parsed,
                        )

                if hasattr(st.session_state, "exam_response"):
                    display_generated_exam(st.session_state.exam_response)
                    # Parse and re-serialise the raw LLM response before download
                    # so the file is guaranteed to be clean, valid JSON regardless
                    # of any fencing or preamble in the original model output.
                    try:
                        _exam_download_data = json.dumps(
                            json.loads(strip_llm_json(st.session_state.exam_response)),
                            indent=2,
                        )
                    except Exception:
                        _exam_download_data = st.session_state.exam_response

                    # Explicit keys are required on both download buttons because
                    # Streamlit renders all tab content simultaneously. Without
                    # explicit keys, auto-generated keys are position-based in the
                    # render tree and can conflict with the equivalent buttons in
                    # Tab 1 which occupy the same structural position.
                    col_dl_json, col_dl_word = st.columns(2)
                    with col_dl_json:
                        st.download_button(
                            label="Download Exam as JSON",
                            data=_exam_download_data,
                            file_name="generated_exam.json",
                            mime="application/json",
                            key="ec_topics_dl_json",
                        )
                    with col_dl_word:
                        try:
                            _exam_word_questions = _parse_from_topics_response(
                                st.session_state.exam_response
                            )
                            # _parse_from_topics_response returns [] silently on any
                            # parse failure. Guard against this so the teacher is
                            # not given a download button that produces an empty
                            # document with no explanation.
                            if not _exam_word_questions:
                                st.warning("Word export is unavailable — the response could not be parsed.")
                            else:
                                _exam_word_buf = _create_exam_word_document(
                                    questions=_exam_word_questions,
                                    creation_mode="from_topics",
                                    title="Generated Exam Questions",
                                )
                                st.download_button(
                                    label="Download Exam as Word",
                                    data=_exam_word_buf.getvalue(),
                                    file_name="generated_exam.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key="ec_topics_dl_word",
                                )
                        except Exception as e:
                            st.error(f"Could not generate Word document: {e}")

    # -------------------------------------------------------------------------
    # TAB 3 — History
    # -------------------------------------------------------------------------
    with tab3:
        st.subheader("Generation History")

        # History is scoped to the current assessment so teachers only see
        # sessions they generated while working in this specific assessment.
        sessions = get_exam_creation_sessions(user_id, assessment_id)

        if not sessions:
            st.info("No generation history for this assessment yet.")
        else:
            for session in sessions:
                session_id     = session["creation_session_id"]
                creation_mode  = session["creation_mode"]
                created_at     = session["created_at"]
                question_count = session["question_count"]

                mode_label = "Variation" if creation_mode == "variation" else "From Topics"
                header     = f"{created_at} — {mode_label} — {question_count} question(s)"

                with st.expander(header, expanded=False):
                    questions = get_exam_creation_session_questions(session_id, user_id)

                    if not questions:
                        st.info("No questions found for this session.")
                    else:
                        for q in questions:
                            q_id     = q["id"]
                            edit_key = f"ec_editing_{q_id}"

                            # Initialise the edit-mode flag for this question if
                            # not already present. False means read-only display.
                            if edit_key not in st.session_state:
                                st.session_state[edit_key] = False

                            # Show the source question for variation-mode sessions
                            # so the teacher knows which original this was derived from.
                            if q.get("original_question_text"):
                                st.caption(f"Original: {q['original_question_text']}")

                            if st.session_state[edit_key]:
                                # ── Edit mode — text areas are active ────────
                                new_text = st.text_area(
                                    "Question",
                                    value=q.get("question_text", ""),
                                    key=f"ec_qtext_{q_id}",
                                )
                                new_guidance = st.text_area(
                                    "Answer Guidance",
                                    value=q.get("answer_guidance", ""),
                                    key=f"ec_guidance_{q_id}",
                                )

                                # Show additional metadata for from-topics questions.
                                if q.get("question_type"):
                                    st.caption(
                                        f"Type: {q['question_type']}  ·  "
                                        f"Topic: {q.get('topic', '—')}  ·  "
                                        f"Difficulty: {q.get('difficulty', '—')}"
                                    )

                                # Save Changes writes to the database and returns
                                # the question to read-only display.
                                if st.button("Save Changes", key=f"ec_save_{q_id}", type="primary"):
                                    update_exam_creation_question(
                                        question_id=q_id,
                                        user_id=user_id,
                                        question_text=new_text,
                                        answer_guidance=new_guidance,
                                    )
                                    st.session_state[edit_key] = False
                                    st.rerun()
                            else:
                                # ── Read-only mode — plain text display ──────
                                st.markdown(f"**Question:** {q.get('question_text', '')}")
                                st.markdown(f"**Answer Guidance:** {q.get('answer_guidance', '') or '—'}")

                                # Show additional metadata for from-topics questions.
                                if q.get("question_type"):
                                    st.caption(
                                        f"Type: {q['question_type']}  ·  "
                                        f"Topic: {q.get('topic', '—')}  ·  "
                                        f"Difficulty: {q.get('difficulty', '—')}"
                                    )

                                # Edit button unlocks the text fields for this question.
                                if st.button("Edit", key=f"ec_edit_{q_id}"):
                                    st.session_state[edit_key] = True
                                    st.rerun()

                            st.divider()

                    # Per-session actions: download as JSON, download as Word, and delete.
                    # Three columns keep the controls in a single row. The questions list
                    # is already fetched above and reflects any edits saved this session,
                    # so both downloads always contain the current DB state.
                    col_dl_json, col_dl_word, col_del = st.columns([1, 1, 1])

                    with col_dl_json:
                        # Build a clean export from the current DB state so
                        # downloads always reflect any edits that have been saved.
                        export = [
                            {
                                "question":        q.get("question_text", ""),
                                "answer_guidance": q.get("answer_guidance", ""),
                                "question_type":   q.get("question_type"),
                                "topic":           q.get("topic"),
                                "difficulty":      q.get("difficulty"),
                            }
                            for q in questions
                        ]
                        st.download_button(
                            label="Download as JSON",
                            data=json.dumps(export, indent=2),
                            file_name=f"exam_{session_id[:8]}.json",
                            mime="application/json",
                            key=f"ec_dl_{session_id}",
                        )

                    with col_dl_word:
                        try:
                            _hist_word_buf = _create_exam_word_document(
                                questions=questions,
                                creation_mode=creation_mode,
                                title=f"Exam Questions — {mode_label}",
                            )
                            st.download_button(
                                label="Download as Word",
                                data=_hist_word_buf.getvalue(),
                                file_name=f"exam_{session_id[:8]}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"ec_dl_word_{session_id}",
                            )
                        except Exception as e:
                            st.error(f"Could not generate Word document: {e}")

                    with col_del:
                        if st.button(
                            "Delete Session",
                            key=f"ec_del_{session_id}",
                            type="primary",
                        ):
                            _dialog_delete_exam_creation_session(session_id, user_id)
