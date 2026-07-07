"""
Quiz Generation Logic

This module contains functions to generate quiz questions using LLMs
based on extracted document content and user preferences.
"""

import json
from typing import Dict, List, Any, Tuple
import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from src.utils.llm_utils import stream_llm, MODELS, MODEL_PROVIDERS, strip_llm_json


def create_quiz_generation_prompt(content: str, question_type: str, num_questions: int, 
                                 difficulty: str, topic_filters: str = "") -> str:
    """
    Create a prompt for the LLM to generate quiz questions.
    
    Args:
        content: The extracted text content from documents
        question_type: Type of questions (Multiple Choice, True/False, Short Answer)
        num_questions: Number of questions to generate
        difficulty: Difficulty level (Easy, Medium, Hard)
        topic_filters: Optional topic filters
        
    Returns:
        str: Formatted prompt for the LLM
    """
    
    # Create the JSON template based on question type
    if question_type == "Multiple Choice":
        json_template = """
{
  "questions": [
    {
      "question_number": 1,
      "question_text": "The text of the question",
      "question_type": "multiple_choice",
      "options": ["Option A", "Option B", "Option C", "Option D"],
      "correct_answer": "Option A",
      "explanation": "Brief explanation of why this is the correct answer",
      "difficulty": "Easy/Medium/Hard",
      "topic": "The topic this question covers"
    }
  ]
}
"""
    elif question_type == "True/False":
        json_template = """
{
  "questions": [
    {
      "question_number": 1,
      "question_text": "The statement to evaluate",
      "question_type": "true_false",
      "options": ["True", "False"],
      "correct_answer": "True",
      "explanation": "Brief explanation of why this is correct",
      "difficulty": "Easy/Medium/Hard",
      "topic": "The topic this question covers"
    }
  ]
}
"""
    else:  # Short Answer
        json_template = """
{
  "questions": [
    {
      "question_number": 1,
      "question_text": "The question requiring a short answer",
      "question_type": "short_answer",
      "options": [],
      "correct_answer": "Expected answer or key points",
      "explanation": "What makes a good answer to this question",
      "difficulty": "Easy/Medium/Hard",
      "topic": "The topic this question covers"
    }
  ]
}
"""
    
    topic_filter_text = f"\nFocus specifically on these topics: {topic_filters}" if topic_filters else ""
    
    prompt = f"""You are an expert educator creating quiz questions based on study materials. 

CONTENT TO ANALYZE:
{content}

REQUIREMENTS:
- Generate {num_questions} questions
- Question type: {question_type}
- Difficulty level: {difficulty}
- Make questions clear, accurate, and educational{topic_filter_text}

INSTRUCTIONS:
1. Carefully read and understand the provided content
2. Create {num_questions} questions that test key concepts from the material
3. For multiple choice questions, ensure all options are plausible but only one is clearly correct
4. For true/false questions, create statements that can be definitively verified from the content
5. For short answer questions, ask for specific information that demonstrates understanding
6. Vary the topics covered to provide comprehensive coverage of the material
7. Match the specified difficulty level:
   - Easy: Direct recall of facts
   - Medium: Application of concepts
   - Hard: Analysis and synthesis of information

OUTPUT FORMAT:
Respond with ONLY valid JSON in this exact format:
{json_template.strip()}

Make sure your response is valid JSON that can be parsed. Include exactly {num_questions} questions.
"""
    
    return prompt


def parse_quiz_response(response: str) -> Dict[str, Any]:
    """
    Parse the LLM response and extract quiz questions.
    
    Args:
        response: Raw response from the LLM
        
    Returns:
        Dict containing parsed quiz questions or error information
    """
    # If the LLM layer returned an error string (e.g. from a Groq rate-limit
    # or network failure), surface it directly rather than attempting to parse
    # it as JSON. Error strings from stream_llm always begin with "Error:".
    # Use the stripped response for this check (not the raw one) so a
    # whitespace-only response - e.g. "\n\n" - is also treated as empty
    # instead of falling through to json.loads() and raising a cryptic
    # "Expecting value: line 1 column 1" error.
    stripped_response = response.strip() if response else ""
    if not stripped_response or stripped_response.startswith("Error:"):
        return {"error": stripped_response or "Empty response from model"}

    try:
        # strip_llm_json extracts the outermost JSON object via bracket-matching,
        # handling bare JSON, fenced JSON, and responses with surrounding prose.
        cleaned = strip_llm_json(response)

        # Some thinking models (e.g. DeepSeek) can get cut off before ever
        # emitting the JSON object, and some responses are pure refusal/prose
        # text that happens to contain a stray "{" (e.g. quoting code or math
        # notation) without ever containing real JSON. strip_llm_json cannot
        # find a balanced object in either case and falls back to returning
        # the raw text untouched, so checking for "{" anywhere in the string
        # isn't enough - a validly extracted candidate always *starts* with
        # "{". Detect the failure here so the error names the actual problem
        # (and shows what the model said) instead of a generic JSON parse
        # failure pointing at line 1 column 1.
        if not cleaned.strip().startswith("{"):
            snippet = response.strip()[:300]
            return {"error": f"Model response did not contain valid JSON. Response started with: {snippet!r}"}

        quiz_data = json.loads(cleaned)

        # Validate the top-level structure.
        if "questions" not in quiz_data:
            return {"error": "Response missing 'questions' field"}

        questions = quiz_data["questions"]
        if not isinstance(questions, list):
            return {"error": "'questions' field must be a list"}

        # Validate each question has the required fields. A response that got
        # cut off mid-generation (e.g. hit the model's output token limit)
        # produces a trailing question missing some fields even after
        # strip_llm_json's truncation repair closes the JSON structurally -
        # drop only the incomplete ones rather than failing the whole batch,
        # so a truncated 10-question response still yields the 8 that
        # finished instead of nothing at all.
        # Check truthiness, not just presence: a truncated/repaired trailing
        # question can end up with e.g. "correct_answer": "" - the key exists
        # but is empty, which validate_quiz_data() (a stricter, separate check
        # downstream in quiz_generator_feature.py) rejects for the *entire*
        # combined multi-type batch. Filtering on emptiness here, in the same
        # place questions are already being dropped for missing fields, keeps
        # the two checks consistent and avoids losing an otherwise-good batch
        # over one incomplete trailing question.
        required_fields = ["question_text", "question_type", "correct_answer"]
        valid_questions = [q for q in questions if all(q.get(f) for f in required_fields)]

        if not valid_questions:
            return {"error": "No complete questions could be parsed from the model response (it may have been cut off before finishing)"}

        quiz_data["questions"] = valid_questions
        return quiz_data

    except json.JSONDecodeError as e:
        # A bare line/column number isn't actionable on its own - show the
        # actual text around the failure point so the real cause (an
        # unescaped character, a missing comma, etc.) is visible immediately
        # instead of requiring another round of guessing from a position number.
        context_start = max(0, e.pos - 80)
        context_end = min(len(e.doc), e.pos + 80)
        context = e.doc[context_start:context_end]
        pointer_offset = e.pos - context_start
        pointer = " " * pointer_offset + "^"
        return {"error": f"Could not parse quiz response: {e.msg} at line {e.lineno} column {e.colno}\n...{context}...\n{' ' * 3}{pointer}"}
    except Exception as e:
        return {"error": f"Could not parse quiz response: {str(e)}"}


def generate_quiz_questions(content: str, question_type: str, num_questions: int,
                          difficulty: str, topic_filters: str, model_id: str) -> Dict[str, Any]:
    """
    Generate quiz questions using the specified LLM.
    
    Args:
        content: Extracted document content
        question_type: Type of questions to generate
        num_questions: Number of questions
        difficulty: Difficulty level
        topic_filters: Topic filters
        model_id: ID of the model to use
        
    Returns:
        Dict containing generated questions or error information
    """
    try:
        # Create the prompt
        prompt = create_quiz_generation_prompt(
            content, question_type, num_questions, difficulty, topic_filters
        )

        # Weaker/local models occasionally ignore the "respond with ONLY
        # valid JSON" instruction entirely and write a plain-text numbered
        # list instead (e.g. "1. Which... - Option A: ...") - no amount of
        # JSON repair can recover that, since there's no JSON to repair. This
        # is usually a one-off rather than a persistent failure, so retry a
        # few times before giving up, same as the proven fix already used in
        # oral_examination_feature.py for the same failure mode. Deliberately
        # NOT using force_json/Ollama's native JSON mode here - that was
        # already tried for this exact quiz-generation use case and found to
        # truncate output to a single question (see the detailed note in
        # oral_examination_feature.py's question-generation call).
        quiz_data = {"error": "No attempts were made"}
        for _attempt in range(3):
            full_response = ""
            for text_chunk in stream_llm(prompt, model_id):
                full_response += text_chunk

            quiz_data = parse_quiz_response(full_response)
            if "error" not in quiz_data:
                break

        if "error" in quiz_data:
            return quiz_data

        # Add metadata
        quiz_data["metadata"] = {
            "question_type": question_type,
            "num_questions": num_questions,
            "difficulty": difficulty,
            "topic_filters": topic_filters,
            "model_used": model_id,
            "content_length": len(content)
        }

        return quiz_data

    except Exception as e:
        return {"error": f"Failed to generate quiz: {str(e)}"}


def validate_quiz_data(quiz_data: Dict[str, Any]) -> bool:
    """
    Validate that the quiz data is properly formatted.
    
    Args:
        quiz_data: Parsed quiz data
        
    Returns:
        bool: True if valid, False otherwise
    """
    if "error" in quiz_data:
        return False
    
    if "questions" not in quiz_data:
        return False
    
    questions = quiz_data["questions"]
    if not isinstance(questions, list) or len(questions) == 0:
        return False
    
    for question in questions:
        if not isinstance(question, dict):
            return False
        
        required_fields = ["question_text", "correct_answer"]
        for field in required_fields:
            if field not in question or not question[field]:
                return False
    
    return True


def generate_multiple_question_types(content: str, mc_count: int, tf_count: int, sa_count: int,
                                     difficulty: str, topic_filters: str, model_id: str) -> Dict[str, Any]:
    """
    Generate quiz questions of multiple types using separate LLM calls.
    
    Args:
        content: Extracted document content
        mc_count: Number of multiple choice questions
        tf_count: Number of true/false questions
        sa_count: Number of short answer questions
        difficulty: Difficulty level
        topic_filters: Topic filters
        model_id: ID of the model to use
        
    Returns:
        Dict containing all generated questions or error information
    """
    all_questions = []
    generation_summary = {
        "multiple_choice": {"requested": mc_count, "generated": 0, "success": False},
        "true_false": {"requested": tf_count, "generated": 0, "success": False},
        "short_answer": {"requested": sa_count, "generated": 0, "success": False}
    }
    
    try:
        # Generate Multiple Choice questions
        if mc_count > 0:
            mc_data = generate_quiz_questions(content, "Multiple Choice", mc_count, difficulty, topic_filters, model_id)
            if "error" not in mc_data and "questions" in mc_data:
                all_questions.extend(mc_data["questions"])
                generation_summary["multiple_choice"]["generated"] = len(mc_data["questions"])
                generation_summary["multiple_choice"]["success"] = True
            else:
                st.warning(f"Failed to generate Multiple Choice questions: {mc_data.get('error', 'Unknown error')}")
        
        # Generate True/False questions
        if tf_count > 0:
            tf_data = generate_quiz_questions(content, "True/False", tf_count, difficulty, topic_filters, model_id)
            if "error" not in tf_data and "questions" in tf_data:
                all_questions.extend(tf_data["questions"])
                generation_summary["true_false"]["generated"] = len(tf_data["questions"])
                generation_summary["true_false"]["success"] = True
            else:
                st.warning(f"Failed to generate True/False questions: {tf_data.get('error', 'Unknown error')}")
        
        # Generate Short Answer questions
        if sa_count > 0:
            sa_data = generate_quiz_questions(content, "Short Answer", sa_count, difficulty, topic_filters, model_id)
            if "error" not in sa_data and "questions" in sa_data:
                all_questions.extend(sa_data["questions"])
                generation_summary["short_answer"]["generated"] = len(sa_data["questions"])
                generation_summary["short_answer"]["success"] = True
            else:
                st.warning(f"Failed to generate Short Answer questions: {sa_data.get('error', 'Unknown error')}")
        
        # Renumber all questions sequentially
        for i, question in enumerate(all_questions):
            question["question_number"] = i + 1
        
        # Create final quiz data
        if all_questions:
            quiz_data = {
                "questions": all_questions,
                "metadata": {
                    "multiple_choice_count": generation_summary["multiple_choice"]["generated"],
                    "true_false_count": generation_summary["true_false"]["generated"],
                    "short_answer_count": generation_summary["short_answer"]["generated"],
                    "total_questions": len(all_questions),
                    "difficulty": difficulty,
                    "topic_filters": topic_filters,
                    "model_used": model_id,
                    "generation_summary": generation_summary
                }
            }
            return quiz_data
        else:
            return {"error": "No questions were successfully generated for any question type"}
    
    except Exception as e:
        return {"error": f"Failed to generate quiz with multiple question types: {str(e)}"}


def create_word_document_questions_only(quiz_data: Dict[str, Any]) -> io.BytesIO:
    """
    Create a Word document containing only the questions (no answers).
    
    Args:
        quiz_data: Dictionary containing quiz questions and metadata
        
    Returns:
        io.BytesIO: Word document as bytes
    """
    doc = Document()
    
    # Add title
    title = doc.add_heading('Practice Exam/Quiz', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add instructions
    doc.add_paragraph('Instructions: Answer all questions to the best of your ability.')
    doc.add_paragraph('')
    
    # Add metadata
    metadata = quiz_data.get("metadata", {})
    if metadata:
        info_para = doc.add_paragraph()
        info_para.add_run(f"Total Questions: {metadata.get('total_questions', 'N/A')}")
        info_para.add_run(f" | Difficulty: {metadata.get('difficulty', 'N/A')}")
        if metadata.get('topic_filters'):
            info_para.add_run(f" | Topics: {metadata.get('topic_filters')}")
        doc.add_paragraph('')
    
    # Add questions
    questions = quiz_data.get("questions", [])
    for question in questions:
        question_num = question.get("question_number", "")
        question_text = question.get("question_text", "")
        question_type = question.get("question_type", "")
        options = question.get("options", [])
        topic = question.get("topic", "")
        
        # Add question
        q_para = doc.add_paragraph()
        q_para.add_run(f"Question {question_num}: ").bold = True
        q_para.add_run(question_text)
        
        if topic and topic != "General":
            q_para.add_run(f" (Topic: {topic})").italic = True
        
        # Add options for multiple choice and true/false
        if question_type in ["multiple_choice", "true_false"] and options:
            for i, option in enumerate(options):
                option_letter = chr(65 + i) if question_type == "multiple_choice" else ""
                option_text = f"{option_letter}) {option}" if option_letter else f"□ {option}"
                doc.add_paragraph(option_text, style='List Bullet')
        
        # Add space for short answer
        elif question_type == "short_answer":
            doc.add_paragraph("Answer:")
            doc.add_paragraph("_" * 60)
            doc.add_paragraph("")
        
        doc.add_paragraph("")  # Add space between questions
    
    # Save to BytesIO
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer


def create_word_document_with_answers(quiz_data: Dict[str, Any]) -> io.BytesIO:
    """
    Create a Word document containing both questions and answers.
    
    Args:
        quiz_data: Dictionary containing quiz questions and metadata
        
    Returns:
        io.BytesIO: Word document as bytes
    """
    doc = Document()
    
    # Add title
    title = doc.add_heading('Practice Exam/Quiz - Answer Key', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    metadata = quiz_data.get("metadata", {})
    if metadata:
        info_para = doc.add_paragraph()
        info_para.add_run(f"Total Questions: {metadata.get('total_questions', 'N/A')}")
        info_para.add_run(f" | Difficulty: {metadata.get('difficulty', 'N/A')}")
        if metadata.get('topic_filters'):
            info_para.add_run(f" | Topics: {metadata.get('topic_filters')}")
        
        # Add generation summary
        gen_summary = metadata.get('generation_summary', {})
        if gen_summary:
            doc.add_paragraph('')
            summary_para = doc.add_paragraph("Question Distribution: ")
            summary_para.add_run(f"Multiple Choice: {gen_summary.get('multiple_choice', {}).get('generated', 0)}, ")
            summary_para.add_run(f"True/False: {gen_summary.get('true_false', {}).get('generated', 0)}, ")
            summary_para.add_run(f"Short Answer: {gen_summary.get('short_answer', {}).get('generated', 0)}")
        
        doc.add_paragraph('')
    
    # Add questions with answers
    questions = quiz_data.get("questions", [])
    for question in questions:
        question_num = question.get("question_number", "")
        question_text = question.get("question_text", "")
        question_type = question.get("question_type", "")
        options = question.get("options", [])
        correct_answer = question.get("correct_answer", "")
        explanation = question.get("explanation", "")
        topic = question.get("topic", "")
        
        # Add question
        q_para = doc.add_paragraph()
        q_para.add_run(f"Question {question_num}: ").bold = True
        q_para.add_run(question_text)
        
        if topic and topic != "General":
            q_para.add_run(f" (Topic: {topic})").italic = True
        
        # Add options for multiple choice and true/false
        if question_type in ["multiple_choice", "true_false"] and options:
            for i, option in enumerate(options):
                option_letter = chr(65 + i) if question_type == "multiple_choice" else ""
                option_text = f"{option_letter}) {option}" if option_letter else f"□ {option}"
                option_para = doc.add_paragraph(option_text, style='List Bullet')
                
                # Highlight correct answer
                if option == correct_answer:
                    option_para.runs[0].bold = True
        
        # Add correct answer section
        answer_para = doc.add_paragraph()
        answer_para.add_run("Correct Answer: ").bold = True
        answer_para.add_run(correct_answer)
        
        if explanation:
            exp_para = doc.add_paragraph()
            exp_para.add_run("Explanation: ").bold = True
            exp_para.add_run(explanation)
        
        doc.add_paragraph("")  # Add space between questions
    
    # Save to BytesIO
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer